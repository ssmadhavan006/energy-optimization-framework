import sys
import re
from pathlib import Path

def build_docx():
    try:
        import docx
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx library is not installed.")
        sys.exit(1)
        
    doc = docx.Document()
    
    # 1. Page Margin Settings (1.0 inch standard margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default style to Times New Roman, 12pt
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # 2. Separate Title Page
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(72)
    title_p.paragraph_format.space_after = Pt(24)
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("EnergyOptAI: An Explainable Machine Learning Framework with Multi-Objective Optimization for Energy-Efficient CNC Machining Parameter Selection")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    author_p = doc.add_paragraph()
    author_p.paragraph_format.space_after = Pt(12)
    author_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run("EnergyOptAI Research Team")
    author_run.font.name = 'Times New Roman'
    author_run.font.size = Pt(12)
    author_run.bold = True
    
    affil_p = doc.add_paragraph()
    affil_p.paragraph_format.space_after = Pt(120)
    affil_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_p.add_run("Department of Mechanical and Industrial Engineering\nCorresponding Email: contact@energyoptai.org")
    affil_run.font.name = 'Times New Roman'
    affil_run.font.size = Pt(11)
    affil_run.italic = True
    
    date_p = doc.add_paragraph()
    date_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("Submission Date: June 28, 2026")
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # 3. Read manuscript.md
    md_path = Path("paper/manuscript.md")
    if not md_path.exists():
        print(f"Error: {md_path} not found.")
        sys.exit(1)
    text = md_path.read_text(encoding="utf-8")
    
    # Clean up markdown headers
    # Strip markdown title block
    text = re.sub(r'# EnergyOptAI: An Explainable Machine[^#]+', '', text)
    
    paragraphs = text.split("\n\n")
    
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
            
        # Parse headings
        if p_text.startswith("# Abstract"):
            # Abstract heading on its own page
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(12)
            heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.paragraph_format.keep_with_next = True
            h_run = heading.add_run("Abstract")
            h_run.font.name = 'Times New Roman'
            h_run.font.size = Pt(14)
            h_run.bold = True
            
        elif p_text.startswith("# "):
            # Main section heading
            title_text = p_text[2:].strip()
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(24)
            heading.paragraph_format.space_after = Pt(12)
            heading.paragraph_format.keep_with_next = True
            h_run = heading.add_run(title_text)
            h_run.font.name = 'Times New Roman'
            h_run.font.size = Pt(14)
            h_run.bold = True
            
        elif p_text.startswith("### ") or p_text.startswith("#### "):
            # Subsections
            title_text = p_text[p_text.find(" "):].strip()
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(16)
            heading.paragraph_format.space_after = Pt(6)
            heading.paragraph_format.keep_with_next = True
            h_run = heading.add_run(title_text)
            h_run.font.name = 'Times New Roman'
            h_run.font.size = Pt(12)
            h_run.bold = True
            
        else:
            # Body content paragraph
            # Handle list formatting
            lines = p_text.splitlines()
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                if line_stripped.startswith("* ") or line_stripped.startswith("- "):
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 2.0  # Double line spacing
                    run = p.add_run(line_stripped[2:])
                    run.font.name = 'Times New Roman'
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 2.0  # Double line spacing
                    p.paragraph_format.first_line_indent = Inches(0.5)  # Indent first line
                    
                    # Convert simple bold/italic formatting
                    # **bold** -> bold
                    cleaned_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line_stripped)
                    # *italic* -> italic
                    cleaned_line = re.sub(r'\*([^*]+)\*', r'\1', cleaned_line)
                    
                    run = p.add_run(cleaned_line)
                    run.font.name = 'Times New Roman'
                    
        # Add abstract page break
        if p_text.startswith("# Abstract"):
            doc.add_page_break()
            
    # Save document
    dest_path = Path("paper/submission_package/manuscript.docx")
    doc.save(dest_path)
    print(f"Generated manuscript.docx at {dest_path}")

if __name__ == "__main__":
    build_docx()
